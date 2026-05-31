from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN = ROOT / "docs" / "PROJECT_REPORT.md"
OUTPUT = ROOT / "docs" / "PROJECT_REPORT.docx"
SCREENSHOT = ROOT / "dashboard_final_edge.png"


def add_code_block(doc: Document, lines: list[str]) -> None:
    if not lines:
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cleaned = []
    for row in rows:
        cells = [cell.strip() for cell in row]
        if all(set(cell) <= {"-"} for cell in cells if cell):
            continue
        cleaned.append(cells)
    if not cleaned:
        return
    table = doc.add_table(rows=len(cleaned), cols=len(cleaned[0]))
    table.style = "Table Grid"
    for row_idx, row in enumerate(cleaned):
        for col_idx, cell in enumerate(row):
            table.cell(row_idx, col_idx).text = cell


def build() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    in_code = False

    for raw_line in MARKDOWN.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                add_table(doc, table_rows)
                table_rows = []
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_rows.append([cell for cell in stripped.strip("|").split("|")])
            continue

        add_table(doc, table_rows)
        table_rows = []

        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped[:2].replace(".", "").isdigit() and ". " in stripped[:5]:
            doc.add_paragraph(stripped.split(". ", 1)[1], style="List Number")
        else:
            doc.add_paragraph(stripped.replace("`", ""))

    add_code_block(doc, code_lines)
    add_table(doc, table_rows)

    if SCREENSHOT.exists():
        doc.add_heading("Dashboard Screenshot", level=1)
        doc.add_paragraph("Captured in Microsoft Edge during final testing.")
        doc.add_picture(str(SCREENSHOT), width=Inches(6.4))

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
